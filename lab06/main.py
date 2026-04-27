import tkinter as tk
from tkinter import ttk, messagebox
from paket import particles, specific_charge, compton_wavelength
from paket.db_postgres import save_result, get_all_results
import docx
from datetime import datetime

class ParticleApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Расчёт частиц")
        self.root.geometry("700x550")
        self.last_particle = self.last_spec = self.last_comp = ""
        
        self.combo = ttk.Combobox(root, values=list(particles.keys()), state="readonly", width=30)
        self.combo.pack(pady=5)
        self.combo.set("Выберите частицу")
        
        self.btn_calc = tk.Button(root, text="Рассчитать", command=self.calculate, bg="lightyellow")
        self.btn_calc.pack(pady=5)
        
        self.result_var = tk.StringVar(value="Нажмите 'Рассчитать'")
        self.result_label = tk.Label(root, textvariable=self.result_var, font=("Arial", 10), justify="left")
        self.result_label.pack(pady=10)
        
        btn_frame = tk.Frame(root)
        btn_frame.pack(pady=5)
        tk.Button(btn_frame, text="Сохранить в БД", command=self.save_to_db, bg="lightblue", width=18).pack(side="left", padx=5)
        tk.Button(btn_frame, text="Сохранить в .docx", command=self.save_doc, bg="lightgreen", width=18).pack(side="left", padx=5)
        
        # Таблица истории
        self.tree = ttk.Treeview(root, columns=("id", "particle", "specific", "compton"), show="headings", height=6)
        for col, width in [("id", 40), ("particle", 100), ("specific", 180), ("compton", 180)]:
            self.tree.heading(col, text=col.capitalize())
            self.tree.column(col, width=width)
        self.tree.pack(pady=10, fill="both", expand=True)
        
        tk.Button(root, text="Обновить историю", command=self.load_history, bg="lightgray").pack(pady=5)
        self.load_history()
    
    def calculate(self):
        particle = self.combo.get()
        if particle == "Выберите частицу":
            messagebox.showwarning("Внимание", "Выберите частицу!")
            return
        
        m, q = particles[particle]["mass"], particles[particle]["charge"]
        spec, comp = specific_charge(q, m), compton_wavelength(m)
        
        self.last_particle, self.last_spec, self.last_comp = particle, f"{spec:.2e}", f"{comp:.2e}"
        
        self.result_var.set(f"{particle}\nУд. заряд: {spec:.2e} Кл/кг\nКомпт. λ: {comp:.2e} м")
        messagebox.showinfo("Успех", "Расчёт выполнен!")
    
    def save_to_db(self):
        if not self.last_particle:
            messagebox.showwarning("Внимание", "Сначала выполните расчёт!")
            return
        try:
            save_result(self.last_particle, self.last_spec, self.last_comp)
            messagebox.showinfo("Успех", "Сохранено в PostgreSQL!")
            self.load_history()
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))
    
    def save_doc(self):
        if not self.last_particle:
            messagebox.showwarning("Внимание", "Сначала выполните расчёт!")
            return
        doc = docx.Document()
        doc.add_heading('Отчёт по частицам', 0)
        doc.add_paragraph(f"Дата: {datetime.now()}\nЧастица: {self.last_particle}\nУдельный заряд: {self.last_spec} Кл/кг\nКомптоновская λ: {self.last_comp} м")
        doc.save(f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        messagebox.showinfo("Готово", "Сохранено в .docx")
    
    def load_history(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        try:
            for row in get_all_results():
                self.tree.insert("", "end", values=(row['id'], row['particle'], row['specific'], row['compton']))
        except:
            self.tree.insert("", "end", values=("", "Нет данных", "", ""))

if __name__ == "__main__":
    root = tk.Tk()           
    app = ParticleApp(root)  
    root.mainloop()          
